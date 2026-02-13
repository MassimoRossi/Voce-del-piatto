import os
import base64
import tempfile
import streamlit as st
from openai import OpenAI
import yaml
import io
import requests
import pandas as pd
from datetime import datetime
from docx import Document

import hmac
import zipfile

def create_archive_zip():
    """Crea un file ZIP contenente l'Excel e tutte le immagini archiviate."""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, 'a', zipfile.ZIP_DEFLATED, False) as zip_file:
        if os.path.exists(ARCHIVE_FILE):
            zip_file.write(ARCHIVE_FILE, ARCHIVE_FILE)
        if os.path.exists(IMAGES_DIR):
            for root, dirs, files in os.walk(IMAGES_DIR):
                for file in files:
                    zip_file.write(os.path.join(root, file), os.path.join(IMAGES_DIR, file))
    return buf.getvalue()

# =====================
# Archiving Logic (Merged from archive_manager.py)
# =====================
ARCHIVE_FILE = "archivio_piatti.xlsx"
IMAGES_DIR = "archived_images"

def initialize_archive():
    """Crea il file Excel con gli header se non esiste."""
    if not os.path.exists(ARCHIVE_FILE):
        df = pd.DataFrame(columns=[
            "seriale", 
            "titolo", 
            "ricetta", 
            "frase_iconica", 
            "immagine_path", 
            "tags", 
            "data_archiviazione"
        ])
        df.to_excel(ARCHIVE_FILE, index=False)
    
    if not os.path.exists(IMAGES_DIR):
        os.makedirs(IMAGES_DIR)

def get_next_serial():
    """Ritorna il prossimo seriale disponibile."""
    if not os.path.exists(ARCHIVE_FILE):
        return 1
    try:
        df = pd.read_excel(ARCHIVE_FILE)
        if df.empty:
            return 1
        return df["seriale"].max() + 1
    except Exception:
        return 1

def add_archive_entry(titolo, ricetta, frase, immagine_bytes, tags):
    """Aggiunge una riga all'archivio Excel e salva l'immagine (se presente)."""
    initialize_archive()
    
    serial = get_next_serial()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    img_path = ""
    if immagine_bytes:
        # Salva immagine usando il seriale come nome file
        img_filename = f"{serial}.png"
        img_path = os.path.join(IMAGES_DIR, img_filename)
        with open(img_path, "wb") as f:
            f.write(immagine_bytes)
    
    # Prepara dati
    new_data = {
        "seriale": serial,
        "titolo": titolo,
        "ricetta": ricetta,
        "frase_iconica": frase,
        "immagine_path": img_filename if immagine_bytes else "",
        "tags": tags,
        "data_archiviazione": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    }
    
    # Carica, appendi e salva
    try:
        # Leggi l'intero foglio per evitare sovrascritture parziali
        if os.path.exists(ARCHIVE_FILE):
            df = pd.read_excel(ARCHIVE_FILE)
        else:
            # Fallback se il file è sparito tra initialize e qui
            df = pd.DataFrame(columns=["seriale", "titolo", "ricetta", "frase_iconica", "immagine_path", "tags", "data_archiviazione"])
        
        # Assicuriamoci che i tipi siano coerenti per il concat
        new_row_df = pd.DataFrame([new_data])
        df = pd.concat([df, new_row_df], ignore_index=True)
        
        # Salvataggio forzato su file
        with pd.ExcelWriter(ARCHIVE_FILE, engine='openpyxl') as writer:
            df.to_excel(writer, index=False)
            
        return serial, img_path
    except Exception as e:
        # Non usare st.error qui se vogliamo che l'errore sia catturato dal chiamante
        raise e

def require_password():
    if st.session_state.get("auth_ok"):
        return

    st.title("Voce del Piatto")
    st.caption("Accesso riservato")

    pwd = st.text_input("Password", type="password")
    if st.button("Entra", use_container_width=True):
        secret = st.secrets.get("APP_PASSWORD", "")
        if secret and hmac.compare_digest(pwd, secret):
            st.session_state["auth_ok"] = True
            st.rerun()
        else:
            st.error("Password non corretta.")

    st.stop()

require_password()

# =====================
# Config
# =====================
st.set_page_config(page_title="Voce del Piatto", layout="wide")  # wide per 3 colonne

client = OpenAI()  # OPENAI_API_KEY da env / Streamlit Secrets

VISION_MODEL = os.getenv("VISION_MODEL", "gpt-4o-mini")
GEN_MODEL = os.getenv("GEN_MODEL", "gpt-4o-mini")
TRANSCRIBE_MODEL = os.getenv("TRANSCRIBE_MODEL", "gpt-4o-transcribe")

# =====================
# Load rules + prompt
# =====================
def load_yaml(path: str):
    with open(path, "r", encoding="utf-8") as f:
        return yaml.safe_load(f)

RULES = load_yaml("rules/registri.yaml")
REGISTRI = RULES["registri"]          # dict: nome -> guida
HARD_RULES = RULES["hard_rules"]      # list[str]

with open("prompts/system.txt", "r", encoding="utf-8") as f:
    SYSTEM_TXT = f.read()

# =====================
# Session state
# =====================
if "ricetta" not in st.session_state:
    st.session_state.ricetta = ""
if "outputs" not in st.session_state:
    st.session_state.outputs = {}  # dict: registro -> testo generato
if "last_params" not in st.session_state:
    st.session_state.last_params = {}

if "draft_text" not in st.session_state:
    st.session_state.draft_text = ""
if "manual_input_text" not in st.session_state:
    st.session_state.manual_input_text = ""
if "confirmed" not in st.session_state:
    st.session_state.confirmed = False
if "recipe_confirmed" not in st.session_state:
    st.session_state.recipe_confirmed = False
if "last_confirmed_ricetta" not in st.session_state:
    st.session_state.last_confirmed_ricetta = ""

if "archival_results" not in st.session_state:
    st.session_state.archival_results = {} # dict: key -> {'excel_bytes': b, 'img_bytes': b, 'serial': s}

# Contatori per resettare i popover
if "pop_counters" not in st.session_state:
    st.session_state.pop_counters = {}

def reset_confirmation():
    st.session_state.recipe_confirmed = False

def clear_all_callback():
    st.session_state.outputs = {}
    st.session_state.ricetta = ""
    st.session_state.manual_input_text = ""
    st.session_state.recipe_confirmed = False

def clear_manual_input_callback():
    st.session_state.manual_input_text = ""

# =====================
# Helpers
# =====================
def _to_data_url(file_bytes: bytes, mime: str) -> str:
    b64 = base64.b64encode(file_bytes).decode("utf-8")
    return f"data:{mime};base64,{b64}"

def extract_text_from_image(image_bytes: bytes, mime: str) -> str:
    data_url = _to_data_url(image_bytes, mime)
    prompt = (
        "Estrai e trascrivi fedelmente il testo della ricetta dall'immagine.\n"
        "Regole:\n"
        "- Non inventare nulla.\n"
        "- Mantieni numeri, unità (g, ml), virgole, simboli e 'q.b.'\n"
        "- Mantieni struttura a righe.\n"
        "- Se c'è una tabella, rendila in testo con colonne separate da ' | '.\n"
        "Output: SOLO il testo estratto."
    )

    resp = client.chat.completions.create(
        model=VISION_MODEL,
        messages=[{
            "role": "user",
            "content": [
                {"type": "text", "text": prompt},
                {"type": "image_url", "image_url": {"url": data_url}},
            ]
        }]
    )
    return (resp.choices[0].message.content or "").strip()

def transcribe_audio_bytes(audio_bytes: bytes, suffix: str) -> str:
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        tmp.write(audio_bytes)
        tmp_path = tmp.name
    try:
        with open(tmp_path, "rb") as f:
            tr = client.audio.transcriptions.create(
                model=TRANSCRIBE_MODEL,
                file=f,
            )
        return getattr(tr, "text", "") or ""
    finally:
        try:
            os.remove(tmp_path)
        except Exception:
            pass

def generate_output(ricetta: str, registro: str, out_type: str, length: str) -> str:
    reg_hint = REGISTRI[registro]

    user_prompt = f"""
Sei un copywriter gastronomico specializzato.
Il tuo compito è scrivere la descrizione di un piatto basandoti sulla ricetta fornita.

COMANDO: Devi generare SOLAMENTE la versione: {out_type} {length}.
NON generare altre varianti.
NON generare introduzioni o spiegazioni.
NON unire più versioni (es. se chiesto Menu, NON fare Cameriere).

REGISTRO RICHIESTO: {registro}
DESCRIZIONE REGISTRO: {reg_hint}
(Usa queste regole di stile SOLO per la versione {out_type} {length})

REGOLE HARD (da rispettare sempre):
- """ + "\n- ".join(HARD_RULES) + f"""

RICETTA (testo sorgente):
{ricetta}

OUTPUT ATTESO:
Scrivi SOLO il testo per {out_type} in formato {length}.
""".strip()

    resp = client.chat.completions.create(
        model=GEN_MODEL,
        messages=[
            {"role": "system", "content": SYSTEM_TXT},
            {"role": "user", "content": user_prompt},
        ],
    )
    return (resp.choices[0].message.content or "").strip()

def translate_text(text: str, language: str, register: str) -> str:
    prompt = f"""
Sei un traduttore esperto di menu gastronomici.
Traduci il seguente testo in {language}.
Mantieni rigorosamente il tono, lo stile e la formattazione del registro originale: "{register}".
Non aggiungere spiegazioni o commenti extra.

TESTO DA TRADURRE:
{text}
""".strip()
    
    resp = client.chat.completions.create(
        model=GEN_MODEL,
        messages=[
            {"role": "user", "content": prompt},
        ],
    )
    return (resp.choices[0].message.content or "").strip()

def export_docx(titolo: str, contenuto: str) -> bytes:
    doc = Document()
    doc.add_heading(titolo, level=1)

    for par in contenuto.split("\n"):
        doc.add_paragraph(par)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

def generate_dish_image(ricetta: str, model="dall-e-3"):
    """Genera l'immagine del piatto in versione ristorante stellato tramite DALL-E."""
    prompt = f"A high-end, gourmet, michelin-starred restaurant version of the following dish: {ricetta}. Professional food photography, elegant plating, soft lighting."
    try:
        response = client.images.generate(
            model=model,
            prompt=prompt,
            size="1024x1024" if model == "dall-e-3" else "512x512",
            quality="standard",
            n=1,
        )
        image_url = response.data[0].url
        img_data = requests.get(image_url).content
        return img_data
    except Exception as e:
        st.error(f"Errore nella generazione dell'immagine ({model}): {e}")
        return None

# =====================
# UX Fine: mini-CSS (pulizia visiva)
# =====================
PRIMARY = "#C9A227"  # zafferano scuro

st.markdown("""
<style>
/* Centra il contenitore principale */
.block-container {
  max-width: 1600px;
  margin-left: auto;
  margin-right: auto;
  padding-top: 1.2rem;
  padding-bottom: 1rem;
}

/* Titolo principale */
h1 {
  color: #C9A227 !important;   /* zafferano */
  text-align: center;
  margin-bottom: 0.2rem;
}

/* Sottotitolo */
[data-testid="stCaptionContainer"] {
  text-align: center;
  margin-top: -0.5rem;
  margin-bottom: 1rem;
}

/* Titoli di sezione */
h2, h3 {
  color: #C9A227;
}

/* Bottoni primari */
button[kind="primary"] {
  background-color: #C9A227 !important;
  border-color: #C9A227 !important;
}
</style>
""", unsafe_allow_html=True)

# =====================
# Header
# =====================
st.markdown("""
<div style="text-align:center; margin-top:2.2rem; margin-bottom:1.2rem;">
  <div style="font-size:3rem; font-weight:800; color:#C9A227; line-height:1.05; letter-spacing:0.5px;">
    Voce<span style="opacity:.9;">•</span>del<span style="opacity:.9;">•</span>Piatto
  </div>
  <div style="font-size:1rem; color:#6b6b6b; margin-top:0.35rem;">
    Il piatto, raccontato bene.
  </div>
</div>
""", unsafe_allow_html=True)

# =====================
# Home / Tool switch
# =====================

page = st.radio(
        "",
        ["Home", "Tool"],
        horizontal=True,
        index=0,
        label_visibility="collapsed"
        )

st.divider()

if page == "Home":
    # --- LANDING ---
    st.markdown("""
    <div style="text-align:center; margin-top:0.5rem; margin-bottom:1.5rem;">
      <div style="font-size:1.6rem; font-weight:700; color:#1F1F1F;">
        Descrizioni per menu e sala, in 30 secondi.
      </div>
      <div style="font-size:1rem; color:#6b6b6b; margin-top:0.4rem;">
        Carica una ricetta (foto o voce), scegli lo stile, genera la tua “voce”.
      </div>
    </div>
    """, unsafe_allow_html=True)

    col1, col2, col3 = st.columns(3, gap="large")

    with col1:
        st.markdown("""
        <div style="background:#F6F6F6; border:1px solid #E6E6E6; padding:14px 16px; border-radius:12px;">
          <div style="font-weight:700; color:#C9A227;">Input naturale</div>
          <div style="margin-top:6px;">Testo, foto o voce → testo editabile.</div>
        </div>
        """, unsafe_allow_html=True)

    with col2:
        st.markdown("""
        <div style="background:#F6F6F6; border:1px solid #E6E6E6; padding:14px 16px; border-radius:12px;">
          <div style="font-weight:700; color:#C9A227;">5 registri</div>
          <div style="margin-top:6px;">Minimal, classico, territoriale, sensoriale, emozionale.</div>
        </div>
        """, unsafe_allow_html=True)

    with col3:
        st.markdown("""
        <div style="background:#F6F6F6; border:1px solid #E6E6E6; padding:14px 16px; border-radius:12px;">
          <div style="font-weight:700; color:#C9A227;">Output pronto</div>
          <div style="margin-top:6px;">Menu o sala, supporto lingue, export DOCX incluso.</div>
        </div>
        """, unsafe_allow_html=True)

    st.stop()

# =====================

# =====================
# Main App Logic
# =====================

reg_names = list(REGISTRI.keys())
default_regs = [r for r in ["Minimal contemporaneo", "Classico elegante"] if r in reg_names]
if not default_regs and reg_names:
    default_regs = [reg_names[0]]
default_idx = reg_names.index("Minimal contemporaneo") if "Minimal contemporaneo" in reg_names else 0

# =====================
# Layout 3 colonne
# =====================
left, center, right = st.columns([2.15, 2.15, 2.15], gap="large")

# =====================
# LEFT: controlli (in form) + azioni
# =====================
with left:
    st.subheader("Impostazioni")

    with st.form("controls", clear_on_submit=False):
        sperimentazione = st.toggle("Confronta stili", value=True, key="sp_sperimentazione")

        if sperimentazione:
            registri_sel = st.multiselect(
            "Registri",
            reg_names,
            default=[],
            key="sp_registri_multi"
    )
        else:
            registro = st.selectbox(
                "Registro",
                reg_names,
                index=default_idx,
                key="sp_registro_single"
            )
            registri_sel = [registro]


        out_type = st.radio("Tipo testo", ["Menu", "Cameriere"], key="sp_out_type")
        length = st.radio("Lunghezza", ["Corto", "Lungo"], key="sp_length")

        # Disable generate if not confirmed
        is_confirmed = st.session_state.get("recipe_confirmed", False)
        
        extra_langs = st.multiselect(
            "Lingue extra",
            ["Inglese", "Tedesco", "Francese", "Spagnolo"],
            default=[],
            key="sp_extra_langs"
        )
        
        genera = st.form_submit_button("Genera", type="primary", disabled=not is_confirmed)


    colA, colB = st.columns(2)
    if colA.button("Pulisci output"):
        st.session_state.outputs = {}
    if colB.button("Pulisci tutto", on_click=clear_all_callback):
        pass


# =====================
# CENTER: input (Foto/Voce/Testo) + revisione
# =====================
with center:
    st.subheader("Input")

    tab_foto, tab_voce, tab_testo = st.tabs(["📷 Foto", "🎙️ Voce", "✍️ Testo"])

    with tab_foto:
        img = st.file_uploader("Carica immagine (JPG/PNG)", type=["jpg", "jpeg", "png"])
        c1, c2 = st.columns([1, 1])
        do_ocr = c1.button("Estrai testo", type="primary")
        if img:
            st.image(img, caption="Anteprima", use_container_width=True)

        if do_ocr:
            if not img:
                st.warning("Carica prima un'immagine.")
            else:
                with st.spinner("Estrazione testo in corso..."):
                    text = extract_text_from_image(img.getvalue(), img.type or "image/jpeg")
                if text:
                    st.session_state.ricetta = text.strip()
                    reset_confirmation()
                    st.success("Testo estratto e copiato in Revisione!")
                    st.rerun()
                else:
                    st.warning("Niente testo utile. Prova foto più nitida/frontale.")


    with tab_voce:
        st.write("MVP: carica una registrazione (wav/mp3/m4a).")
        aud = st.file_uploader("Carica audio", type=["wav", "mp3", "m4a", "aac"])
        do_stt = st.button("Trascrivi audio", type="primary")
        if aud:
            st.audio(aud)

        if do_stt:
            if not aud:
                st.warning("Carica prima un file audio.")
            else:
                suffix = os.path.splitext(aud.name)[1] or ".wav"
                with st.spinner("Trascrizione in corso..."):
                    text = transcribe_audio_bytes(aud.getvalue(), suffix=suffix)
                if text:
                    st.session_state.ricetta = text.strip()
                    reset_confirmation()
                    st.success("Trascrizione completata e copiata in Revisione!")
                    st.rerun()
                else:
                    st.warning("Trascrizione vuota. Prova un audio più pulito.")

    with tab_testo:
        txt_in = st.text_area("Scrivi/Incolla", key="manual_input_text", height=200)

        col_t1, col_t2 = st.columns(2)
        if col_t1.button("Copia in Revisione", type="primary", use_container_width=True):
            if txt_in and txt_in.strip():
                st.session_state.ricetta = txt_in.strip()
                reset_confirmation()
                st.rerun()
            else:
                st.warning("Scrivi qualcosa prima di copiare.")
        
        if col_t2.button("Cancella testo", use_container_width=True, on_click=clear_manual_input_callback):
            pass


    st.divider()

    st.subheader("Revisione ricetta")
    val_ricetta = st.text_area(
            "Testo ufficiale (usato per generare)",
            key="ricetta",
            height=260,
            on_change=reset_confirmation
    )

    # Confirmation button
    if st.button("✅ Conferma Ricetta", type="primary", use_container_width=True):
        if val_ricetta.strip():
            st.session_state.recipe_confirmed = True
            st.session_state.last_confirmed_ricetta = val_ricetta
            st.rerun()
        else:
            st.warning("La ricetta è vuota.")
    
    if st.session_state.get("recipe_confirmed", False):
        st.success("Ricetta confermata. Puoi generare.")



# =====================
# RIGHT: output persistente
# =====================
with right:
    st.subheader("Output")

    ricetta = (st.session_state.ricetta or "").strip()

    if genera:
        # SICUREZZA: Verifica se il testo è cambiato rispetto all'ultima conferma
        if st.session_state.get("recipe_confirmed") and ricetta != st.session_state.get("last_confirmed_ricetta", ""):
            st.session_state.recipe_confirmed = False
            st.error("⚠️ Hai modificato il testo: devi confermare di nuovo la ricetta.")
            st.stop()

        if not ricetta:
             # Should be caught by disabled, but safety check
            st.error("Manca la ricetta.")
        elif not st.session_state.get("recipe_confirmed", False):
             st.error("Devi confermare la ricetta.")
        else:
            st.session_state.outputs = {}
            st.session_state.last_params = {
                "registri": registri_sel,
                "tipo": out_type,
                "lunghezza": length
            }
            for r in registri_sel:
                with st.spinner(f"Genero: {r}…"):
                    # 1. Italian generation
                    base_text = generate_output(ricetta, r, out_type, length)
                    final_output = base_text
                    
                    # 2. Translate if needed
                    for lang in extra_langs:
                         with st.spinner(f"Traduco in {lang}..."):
                             tr_text = translate_text(base_text, lang, r)
                             final_output += f"\n\n--- {lang.upper()} ---\n{tr_text}"
                    
                    st.session_state.outputs[r] = final_output

    # Mostra sempre ultimo output generato (persistente)
    if st.session_state.outputs:
        params = st.session_state.last_params or {}
        st.caption(f"Ultima generazione: {params.get('tipo','')} / {params.get('lunghezza','')}")

        for r, txt in st.session_state.outputs.items():
            with st.expander(r, expanded=True):
                st.write(txt)
                st.code(txt, language="markdown")

                filename = f"{r}_{params.get('tipo','')}_{params.get('lunghezza','')}.docx".replace(" ", "_")
                titolo = f"Voce del Piatto — {r} — {params.get('tipo','')} — {params.get('lunghezza','')}"
                docx_bytes = export_docx(titolo, txt)
                # --- AZIONI (Download e Archiviazione) ---
                col_d1, col_d2 = st.columns(2)
                
                with col_d1:
                    st.download_button(
                        "Scarica DOCX",
                        data=docx_bytes,
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"btn_dl_{r.replace(' ', '_')}",
                        use_container_width=True
                    )

                with col_d2:
                    # Usiamo un contatore per forzare la chiusura del popover quando necessario
                    pop_key_id = f"pop_{r.replace(' ', '_')}"
                    if pop_key_id not in st.session_state.pop_counters:
                        st.session_state.pop_counters[pop_key_id] = 0
                    
                    # Usiamo un popover per raccogliere titolo e tags in modo pulito
                    with st.popover("Archivia piatto", use_container_width=True, key=f"{pop_key_id}_{st.session_state.pop_counters[pop_key_id]}"):
                        st.subheader("Dati per l'archivio:")
                        
                        # --- SINCRONIZZAZIONE INTEGRATA ---
                        with st.expander("🔄 Sincronizza archivio locale", expanded=not os.path.exists(ARCHIVE_FILE)):
                            up_file = st.file_uploader("Carica il tuo Excel per non perdere le modifiche", type=["xlsx"], key=f"up_{r.replace(' ', '_')}")
                            if up_file:
                                file_key = f"{up_file.name}_{up_file.size}"
                                if st.session_state.get("last_synced_file") != file_key:
                                    with open(ARCHIVE_FILE, "wb") as f:
                                        f.write(up_file.getbuffer())
                                    st.session_state["last_synced_file"] = file_key
                                    st.success("Archivio sincronizzato!")
                                    st.rerun()
                        
                        titolo_piatto = st.text_input("Titolo del piatto", key=f"title_{r.replace(' ', '_')}")
                        tags_piatto = st.text_input("Tags (es. mare, primo, etc.)", key=f"tags_{r.replace(' ', '_')}")
                        
                        st.divider()
                        do_gen_img = st.checkbox("Genera immagine AI", value=True, key=f"gen_img_{r.replace(' ', '_')}")
                        img_model = st.selectbox("Modello immagine", ["dall-e-3", "dall-e-2"], index=0, key=f"model_img_{r.replace(' ', '_')}", disabled=not do_gen_img)
                        
                        if st.button("Conferma Archiviazione", key=f"btn_arch_{r.replace(' ', '_')}", type="primary", use_container_width=True):
                            if not titolo_piatto.strip():
                                st.warning("Inserisci un titolo per il piatto.")
                            else:
                                with st.spinner("Archiviazione in corso..."):
                                    try:
                                        img_bytes = None
                                        if do_gen_img:
                                            with st.spinner(f"Generazione immagine ({img_model})..."):
                                                img_bytes = generate_dish_image(ricetta, model=img_model)
                                        
                                        serial, img_filename = add_archive_entry(
                                            titolo=titolo_piatto.strip(),
                                            ricetta=ricetta,
                                            frase=txt,
                                            immagine_bytes=img_bytes,
                                            tags=tags_piatto.strip()
                                        )
                                        
                                        if serial:
                                            # Rilegge per avere la versione aggiornata
                                            with open(ARCHIVE_FILE, "rb") as f:
                                                excel_bytes = f.read()
                                            
                                            st.session_state.archival_results[f"res_{r.replace(' ', '_')}"] = {
                                                "excel_bytes": excel_bytes,
                                                "img_bytes": img_bytes,
                                                "serial": serial,
                                                "img_filename": img_filename
                                            }
                                            st.success(f"Piatto archiviato! (Seriale: {serial})")
                                            st.balloons()
                                        else:
                                            st.error("Errore durante il salvataggio dei dati.")
                                    except Exception as e:
                                        st.error(f"Errore durante l'archiviazione: {e}")
                        
                        # --- DOWNLOAD IMMEDIATI (DOPO SUCCESSO) ---
                        res_key = f"res_{r.replace(' ', '_')}"
                        if res_key in st.session_state.archival_results:
                            res = st.session_state.archival_results[res_key]
                            st.divider()
                            st.write("📥 Scarica subito sul tuo PC:")
                            
                            st.download_button(
                                "📊 Scarica Excel Aggiornato",
                                data=res["excel_bytes"],
                                file_name="archivio_piatti.xlsx",
                                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                key=f"dl_xl_{res_key}",
                                use_container_width=True
                            )
                            
                            if res["img_bytes"]:
                                st.download_button(
                                    f"🖼️ Scarica Immagine ({res['img_filename']})",
                                    data=res["img_bytes"],
                                    file_name=res["img_filename"],
                                    mime="image/png",
                                    key=f"dl_img_{res_key}",
                                    use_container_width=True
                                )
                        
                        if st.button("Annulla / Chiudi", key=f"btn_canc_{r.replace(' ', '_')}", use_container_width=True):
                            # Incrementiamo il contatore per forzare la chiusura del popover (cambiando la sua key)
                            st.session_state.pop_counters[pop_key_id] += 1
                            # Pulisce anche i risultati precedenti
                            if res_key in st.session_state.archival_results:
                                del st.session_state.archival_results[res_key]
                            st.rerun()

    else:
        st.markdown("""
        <div style="
            background:#F6F6F6;
            border:1px solid #E6E6E6;
            padding:14px 16px;
            border-radius:10px;
            color:#1F1F1F;">
            <b>Nessun output ancora.</b> Inserisci ricetta e premi <b>Genera</b>.
        </div>
        """, unsafe_allow_html=True)
